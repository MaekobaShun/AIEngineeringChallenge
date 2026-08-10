"""python-docx based parser: paragraphs/headings/tables with run-level
formatting (bold/italic/underline/strike/highlight/font color).

Known gap: .docx has no stored pagination (it's reflowable; page breaks
depend on the renderer/printer), so we cannot report real page numbers from
the file alone. Questions that ask "what page is heading X on" need either
an actual layout engine (Word/LibreOffice rendering to PDF) or must be
answered from a page-numbered format instead (PDF). Not implemented yet --
tracked as an open gap, see README.
"""

from __future__ import annotations

import docx
from docx.table import Table

from datastel_rag.ingest.models import Block, FormattedRun, ParsedDocument


def _run_format(run) -> FormattedRun:
    color = None
    try:
        if run.font.color and run.font.color.rgb:
            color = str(run.font.color.rgb)
    except Exception:
        pass
    highlight = None
    try:
        if run.font.highlight_color is not None:
            highlight = run.font.highlight_color.name.lower()
    except Exception:
        pass
    return FormattedRun(
        text=run.text,
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
        strike=bool(run.font.strike) if run.font.strike is not None else False,
        font_color=color,
        highlight=highlight,
    )


def _paragraph_block(para, block_id: str) -> Block | None:
    if not para.text.strip():
        return None
    style_name = (para.style.name if para.style else "") or ""
    kind = "heading" if style_name.lower().startswith("heading") or style_name.lower() == "title" else "paragraph"
    runs = [_run_format(r) for r in para.runs if r.text]
    return Block(
        block_id=block_id,
        kind=kind,
        text=para.text,
        runs=runs,
        location={"style": style_name} if style_name else {},
    )


def _table_blocks(table: Table, block_id_prefix: str) -> list[Block]:
    blocks = []
    for ri, row in enumerate(table.rows):
        cell_texts = []
        for cell in row.cells:
            cell_texts.append(cell.text)
        text = " | ".join(cell_texts)
        if not text.strip():
            continue
        blocks.append(
            Block(
                block_id=f"{block_id_prefix}/row{ri}",
                kind="table_row",
                text=text,
                location={"row": ri},
            )
        )
    return blocks


def _iter_body_elements(document):
    """Yield ("paragraph", obj) / ("table", obj) in document order, including
    tables nested at top level (nested-in-cell tables are walked separately
    via table.rows -> cell.tables if ever needed)."""
    from docx.oxml.ns import qn

    body = document.element.body
    para_iter = iter(document.paragraphs)
    table_iter = iter(document.tables)
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("paragraph", next(para_iter))
        elif child.tag == qn("w:tbl"):
            yield ("table", next(table_iter))


def parse_docx(abs_path: str, source_rel_path: str, project_key: str) -> ParsedDocument:
    doc = ParsedDocument(source_rel_path=source_rel_path, project_key=project_key, ext="docx")
    try:
        d = docx.Document(abs_path)
    except Exception as e:
        doc.parse_errors.append(f"open failed: {e}")
        return doc

    idx = 0
    table_idx = 0
    for kind, obj in _iter_body_elements(d):
        if kind == "paragraph":
            b = _paragraph_block(obj, f"para{idx}")
            if b:
                doc.blocks.append(b)
            idx += 1
        else:
            doc.blocks.extend(_table_blocks(obj, f"table{table_idx}"))
            table_idx += 1

    core = d.core_properties
    doc.meta = {
        "author": core.author or "",
        "last_modified_by": core.last_modified_by or "",
        "created": str(core.created) if core.created else "",
        "modified": str(core.modified) if core.modified else "",
        "num_paragraphs": len(d.paragraphs),
        "num_tables": len(d.tables),
    }
    return doc
